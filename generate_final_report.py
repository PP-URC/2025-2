"""
generate_final_report.py
-------------------------------------
1. Loads unrc.db and aggregates data
2. Fits logistic regression
3. Creates figures (0a–3) in Spanish
4. Embeds them into the Word report with references
5. Saves PP_UNRC_reporte_final.docx
6. Prints summaries to console
7. Appends Anexo with text tables
"""

import os
import sqlite3
import pandas as pd
import matplotlib.pyplot as plt
from statsmodels.discrete.discrete_model import Logit
from statsmodels.tools import add_constant
from docx import Document
from docx.shared import Inches

# --- Paths ---
DB_PATH = "unrc.db"
OUT_DIR = "./out_pipeline"
REPORT_BASE = "PP_UNRC_reporte_refined.docx"
REPORT_FINAL = "PP_UNRC_reporte_final.docx"

os.makedirs(OUT_DIR, exist_ok=True)

# --- 1. Load data ---
conn = sqlite3.connect(DB_PATH)
students = pd.read_sql("SELECT * FROM students_raw", conn)
panel = pd.read_sql("SELECT * FROM inscripciones", conn)

# Derive abandono if not exists
if "abandono" not in panel.columns:
    panel["abandono"] = 0
if "stopout" not in panel.columns:
    panel["stopout"] = 0

print(panel.columns.tolist())
print(panel.head())


# --- 2. Aggregate ---
agg_sem = panel.groupby("semestre").agg(
    abandono_sem=("abandono","mean"),
    stopout_sem=("stopout","mean"),
    n=("student_id","count")
).reset_index()

# --- 3. Logistic regression ---
X = panel[["promedio","asistencia_pct","horas_trabajo","traslado_min"]].copy()
X = add_constant(X, has_constant="add")
y = panel["abandono"]
logit = Logit(y, X).fit(disp=False)

# --- Print summaries to console ---
print("\n📊 Ejemplo de datos crudos (students_raw):")
print(students.head(10))
print("\n📊 Tabla agregada por semestre (agg_sem):")
print(agg_sem.head(10))
print("\n📊 Resumen de regresión logística:")
print(logit.summary())

# --- 4. Figures ---

# Figura 0a: raw sample
fig, ax = plt.subplots(figsize=(8,2))
ax.axis("off")
tbl = ax.table(cellText=students.head(10).values,
               colLabels=students.columns,
               cellLoc="center", loc="center")
tbl.auto_set_font_size(False); tbl.set_fontsize(8); tbl.scale(1.2, 1.2)
plt.title("Ejemplo de datos crudos (students_raw)", fontsize=10)
plt.savefig(os.path.join(OUT_DIR, "figura0a_datos_crudos.png"), bbox_inches="tight")
plt.close()

# Figura 0b: aggregated sample
fig, ax = plt.subplots(figsize=(6,2))
ax.axis("off")
tbl = ax.table(cellText=agg_sem.head(10).values,
               colLabels=agg_sem.columns,
               cellLoc="center", loc="center")
tbl.auto_set_font_size(False); tbl.set_fontsize(8); tbl.scale(1.2, 1.2)
plt.title("Ejemplo de tabla agregada por semestre (agg_sem)", fontsize=10)
plt.savefig(os.path.join(OUT_DIR, "figura0b_datos_agregados.png"), bbox_inches="tight")
plt.close()

# Figura 1: abandono por semestre
plt.figure()
agg_sem.plot(x="semestre", y="abandono_sem", marker="o", legend=False)
plt.title("Tasa de abandono por semestre")
plt.ylabel("Proporción de abandono")
plt.xlabel("Semestre")
plt.grid(True)
plt.savefig(os.path.join(OUT_DIR, "figura1_abandono_por_semestre.png"))
plt.close()

# Figura 2: stopout por semestre
plt.figure()
agg_sem.plot(x="semestre", y="stopout_sem", marker="o", color="orange", legend=False)
plt.title("Tasa de reingreso temporal (stop-out) por semestre")
plt.ylabel("Proporción de stop-out")
plt.xlabel("Semestre")
plt.grid(True)
plt.savefig(os.path.join(OUT_DIR, "figura2_stopout_por_semestre.png"))
plt.close()

# Figura 3: logistic regression coefficients
coefs = pd.DataFrame({
    "var": X.columns,
    "coef": logit.params,
    "pval": logit.pvalues
})
coefs = coefs[coefs["var"] != "const"].sort_values("coef")

plt.figure(figsize=(6,4))
plt.barh(coefs["var"], coefs["coef"], color="steelblue")
plt.title("Coeficientes de la regresión logística para abandono")
plt.xlabel("Efecto en log-odds de abandono")
plt.tight_layout()
plt.savefig(os.path.join(OUT_DIR, "figura3_coef_logistica.png"))
plt.close()

# --- 5. Update Report ---
doc = Document(REPORT_BASE)

# Add references to text
for para in doc.paragraphs:
    if "Se construyó una base SQLite con dos tablas" in para.text:
        para.text += " (véanse Figuras 0a y 0b para ejemplos de los datos crudos y de la tabla agregada)."
    if "Los mayores niveles de abandono ocurrieron" in para.text:
        para.text += " (véase Figura 1)."
    if "Variables críticas:" in para.text:
        para.text = para.text.replace("Variables críticas:", "Variables críticas (véase Figura 2 para stop-outs):")
    if "coeficientes de regresión" in para.text.lower():
        para.text += " (véase Figura 3)."

# Insert Fig 0a, 0b before section 3
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("3. Metodología"):
        insert_index = i
        break
doc.paragraphs[insert_index].insert_paragraph_before("Figura 0b. Fragmento de tabla agregada por semestre (agg_sem).")
doc.paragraphs[insert_index].insert_paragraph_before("").add_run().add_picture(
    f"{OUT_DIR}/figura0b_datos_agregados.png", width=Inches(5.5))
doc.paragraphs[insert_index].insert_paragraph_before("Figura 0a. Fragmento de datos crudos de estudiantes (students_raw).")
doc.paragraphs[insert_index].insert_paragraph_before("").add_run().add_picture(
    f"{OUT_DIR}/figura0a_datos_crudos.png", width=Inches(5.5))

# Insert Fig 1 & 2 after Resultados heading
for i, para in enumerate(doc.paragraphs):
    if "4. Resultados Principales" in para.text:
        insert_index = i+1
        break
doc.paragraphs[insert_index].insert_paragraph_before("Figura 2. Tasa de reingreso temporal (stop-out) por semestre (datos sintéticos).")
doc.paragraphs[insert_index].insert_paragraph_before("").add_run().add_picture(
    f"{OUT_DIR}/figura2_stopout_por_semestre.png", width=Inches(5.5))
doc.paragraphs[insert_index].insert_paragraph_before("Figura 1. Tasa de abandono por semestre (datos sintéticos).")
doc.paragraphs[insert_index].insert_paragraph_before("").add_run().add_picture(
    f"{OUT_DIR}/figura1_abandono_por_semestre.png", width=Inches(5.5))

# Insert Fig 3 after coefficients discussion
for j, para in enumerate(doc.paragraphs):
    if "coeficientes de regresión" in para.text.lower():
        insert_index_2 = j+1
        break
doc.paragraphs[insert_index_2].insert_paragraph_before("Figura 3. Coeficientes de la regresión logística para abandono.")
doc.paragraphs[insert_index_2].insert_paragraph_before("").add_run().add_picture(
    f"{OUT_DIR}/figura3_coef_logistica.png", width=Inches(5.5))

# --- 6. Add Anexo with text tables ---
doc.add_page_break()
doc.add_heading("Anexo: Tablas de ejemplo", level=1)

doc.add_heading("Tabla A1. Fragmento de datos crudos (students_raw)", level=2)
t = doc.add_table(rows=1, cols=len(students.columns))
hdr_cells = t.rows[0].cells
for i, col in enumerate(students.columns):
    hdr_cells[i].text = col
for _, row in students.head(10).iterrows():
    cells = t.add_row().cells
    for i, col in enumerate(students.columns):
        cells[i].text = str(row[col])

doc.add_paragraph("")

doc.add_heading("Tabla A2. Fragmento de tabla agregada (agg_sem)", level=2)
t = doc.add_table(rows=1, cols=len(agg_sem.columns))
hdr_cells = t.rows[0].cells
for i, col in enumerate(agg_sem.columns):
    hdr_cells[i].text = col
for _, row in agg_sem.head(10).iterrows():
    cells = t.add_row().cells
    for i, col in enumerate(agg_sem.columns):
        cells[i].text = str(row[col])

# --- Save final report ---
doc.save(REPORT_FINAL)
print(f"\n✅ Reporte final generado: {REPORT_FINAL}")
